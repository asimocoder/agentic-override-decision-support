"""
docx_generator.py
Generates a formatted DOCX intelligence brief from the Researcher output.

Called by uw_helper.py after run_agent2_step completes.
Also called by pdf_generator.py as an intermediate step before LibreOffice conversion.
"""

import re
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette (matches a clean professional report) ─────────────────────
DARK_BLUE = RGBColor(0x1F, 0x3A, 0x5F)   # headings, header text
NEAR_BLACK = RGBColor(0x1A, 0x1A, 0x1A)  # body text
MID_GREY = RGBColor(0x5A, 0x5A, 0x5A)    # metadata lines
DIVIDER_GREY = RGBColor(0xCC, 0xCC, 0xCC)


# ── Low-level helpers ─────────────────────────────────────────────────────────

def _set_run(run, text: str, bold: bool = False, italic: bool = False,
             size_pt: int = 11, colour: RGBColor = NEAR_BLACK):
    run.text = text
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    run.font.color.rgb = colour
    run.font.name = "Calibri"


def _add_paragraph(doc: Document, text: str = "", bold: bool = False,
                   italic: bool = False, size_pt: int = 11,
                   colour: RGBColor = NEAR_BLACK,
                   alignment=WD_ALIGN_PARAGRAPH.LEFT,
                   space_before: int = 0, space_after: int = 4) -> None:
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    if text:
        run = para.add_run()
        _set_run(run, text, bold=bold, italic=italic,
                 size_pt=size_pt, colour=colour)


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin grey bottom border to a paragraph to simulate a divider."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(6)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_metadata_row(doc: Document, label: str, value: str) -> None:
    """Single-line metadata row: 'Label:  Value' with label in grey, value in near-black."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(2)
    label_run = para.add_run(f"{label}:   ")
    _set_run(label_run, f"{label}:   ", size_pt=10, colour=MID_GREY)
    value_run = para.add_run(value)
    _set_run(value_run, value, size_pt=10, colour=NEAR_BLACK)


# ── Brief text parser ─────────────────────────────────────────────────────────

def _parse_brief_sections(brief_text: str) -> list[tuple[str, str]]:
    """
    Parse the brief into (section_heading, body) tuples.
    Sections are detected by lines that are ALL CAPS or match known heading patterns.
    Falls back to rendering the brief as a single block if no sections detected.
    Returns list of (heading, body) where heading may be "" for the opening block.
    """
    # Strip the RECOMMENDED STANCE line — already shown in header
    lines = brief_text.strip().splitlines()
    cleaned_lines = [
        l for l in lines
        if not l.strip().upper().startswith("RECOMMENDED STANCE:")
    ]

    # Known section heading patterns from the Researcher system prompt
    heading_patterns = [
        r"^KEY POSITIVE SIGNALS",
        r"^KEY RISK SIGNALS",
        r"^RESOLVED AND UNRESOLVED QUERIES",
        r"^RESOLVED QUERIES",
        r"^UNRESOLVED QUERIES",
        r"^RECOMMENDED STANCE",
        r"^RATIONALE",
        r"^SUMMARY",
        r"^INTELLIGENCE BRIEF",
        r"^RESEARCH FINDINGS",
        r"^DIRECTOR",
        r"^RELATED",
        r"^PRIMARY COMPANY",
    ]

    sections = []
    current_heading = ""
    current_body_lines = []

    for line in cleaned_lines:
        stripped = line.strip()
        clean = _strip_markdown(stripped)
        is_heading = any(
            re.match(pat, clean.upper())
            for pat in heading_patterns
        ) or (
            clean.isupper()
            and len(stripped) > 3
            and len(stripped) < 80
            and not stripped.startswith("-")
        )

        if is_heading and clean:
            if current_body_lines or current_heading:
                sections.append((current_heading, "\n".join(current_body_lines).strip()))
            current_heading = clean.title()
            current_body_lines = []
        else:
            current_body_lines.append(line)

    if current_body_lines or current_heading:
        sections.append((current_heading, "\n".join(current_body_lines).strip()))

    # If no sections detected, return as single block
    if not sections or all(h == "" for h, _ in sections):
        return [("", brief_text.strip())]

    return sections


def _strip_markdown(text: str) -> str:
    """Strip common markdown from a line -- bold, italic, heading hashes."""
    # Remove heading hashes
    text = re.sub(r'^#{1,6}\s*', '', text)
    # Remove bold/italic -- preserve inner text
    text = re.sub(r'\*{1,3}(.*?)\*{1,3}', lambda m: m.group(1), text)
    # Remove inline code -- preserve inner text
    text = re.sub(r'`(.*?)`', lambda m: m.group(1), text)
    return text.strip()

def _render_brief_body(doc: Document, brief_text: str) -> None:
    """
    Render the brief text with section headings and body paragraphs.
    Bullet lines (starting with - or •) are rendered with a small indent.
    """
    sections = _parse_brief_sections(brief_text)

    for heading, body in sections:
        if heading:
            _add_paragraph(
                doc, heading,
                bold=True, size_pt=11,
                colour=DARK_BLUE,
                space_before=10, space_after=3
            )

        if not body:
            continue

        for line in body.splitlines():
            stripped = line.strip()
            if not stripped:
                # Blank line — small spacing gap
                _add_paragraph(doc, "", size_pt=4, space_before=0, space_after=2)
                continue

            is_bullet = stripped.startswith(("-", "•", "*", "–"))

            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(3)

            if is_bullet:
                para.paragraph_format.left_indent = Inches(0.25)
                text = _strip_markdown(stripped.lstrip("-•*– ").strip())
            else:
                text = _strip_markdown(stripped)

            run = para.add_run(text)
            _set_run(run, text, size_pt=10.5, colour=NEAR_BLACK)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_docx(
    brief: str,
    go_nogo: str,
    company_name: str,
    trigger_type: str,
    output_path: str | None = None,
) -> str:
    """
    Generate a DOCX intelligence brief.

    Args:
        brief:        Full brief text from Researcher agent
        go_nogo:      Stance string — "GO", "NOGO", or "NEEDS FURTHER RESEARCH"
        company_name: Borrower company name (from Step 3 input)
        trigger_type: Trigger type selected in Step 1
        output_path:  Optional explicit output path. If None, writes to a temp file.

    Returns:
        Path to the generated DOCX file as a string.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    # ── Header block ──────────────────────────────────────────────────────────
    # System name — small, dark blue, not dominant
    _add_paragraph(
        doc,
        "Agentic Override Decision Support System",
        bold=False, size_pt=9,
        colour=DARK_BLUE,
        space_before=0, space_after=1
    )

    # Title line
    _add_paragraph(
        doc,
        "MSME Credit Intelligence Brief",
        bold=True, size_pt=13,
        colour=DARK_BLUE,
        space_before=2, space_after=6
    )

    # Metadata rows
    date_str = datetime.now().strftime("%d %B %Y")
    _add_metadata_row(doc, "Date", date_str)
    _add_metadata_row(doc, "Company", company_name)
    _add_metadata_row(doc, "Trigger Type", trigger_type)

    # Recommended stance — slightly bold, same size as body, not oversized
    stance_para = doc.add_paragraph()
    stance_para.paragraph_format.space_before = Pt(4)
    stance_para.paragraph_format.space_after = Pt(4)
    label_run = stance_para.add_run("Recommended Stance:   ")
    _set_run(label_run, "Recommended Stance:   ", size_pt=10, colour=MID_GREY)
    stance_run = stance_para.add_run(go_nogo)
    _set_run(stance_run, go_nogo, bold=True, size_pt=10.5, colour=NEAR_BLACK)

    # Divider
    _add_horizontal_rule(doc)

    # ── Brief body ────────────────────────────────────────────────────────────
    _render_brief_body(doc, brief)

    # ── Save ──────────────────────────────────────────────────────────────────
    if output_path is None:
        tmp = tempfile.NamedTemporaryFile(
            suffix=".docx", delete=False,
            prefix=f"brief_{company_name.replace(' ', '_')}_"
        )
        output_path = tmp.name
        tmp.close()

    doc.save(output_path)
    return output_path
